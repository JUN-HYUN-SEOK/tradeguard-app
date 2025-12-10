import streamlit as st
import pandas as pd
import numpy as np
import os
from docx import Document
import datetime
import traceback
import io
import time
import plotly.express as px
import plotly.graph_objects as go

# --- Constants ---
COL_TARIFF_RATE = '관세실행세율'
COL_RATE_TYPE = '세율구분'
COL_TAX_CLASSIFICATION = '세율구분'  # COL_RATE_TYPE과 동일
COL_IMPORT_DEC_NO = '수입신고번호'
COL_ACCEPTANCE_DATE = '수리일자'
COL_BL_NO = 'B/L번호'
COL_HS_CODE = '세번부호'
COL_RATE_DESC = '세율설명'
COL_EXPORT_COUNTRY = '적출국코드'
COL_ORIGIN_COUNTRY = '원산지코드'
COL_SPEC_1 = '규격1'
COL_SPEC_2 = '규격2'
COL_SPEC_3 = '규격3'
COL_COMP_1 = '성분1'
COL_COMP_2 = '성분2'
COL_COMP_3 = '성분3'
COL_ACTUAL_DUTY = '실제관세액'
COL_PAYMENT_METHOD = '결제방법'
COL_CURRENCY = '결제통화단위'
COL_TRADE_COMPANY = '무역거래처상호'
COL_TRADE_COUNTRY = '무역거래처국가코드'
COL_TRADE_NAME = '거래품명'
COL_LINE_NO = '란번호'
COL_ROW_NO = '행번호'
COL_QTY_1 = '수량_1'
COL_UNIT_1 = '수량단위_1'
COL_UNIT_PRICE = '단가'
COL_AMOUNT = '금액'
COL_LINE_PAYMENT_AMT = '란결제금액'
COL_ROW_DUTY = '행별관세'
COL_FTA_REVIEW = 'FTA사후환급 검토'
COL_TRADE_TYPE = '거래구분'
COL_INTERNAL_TAX_CODE = '내국세부호'
COL_TAXABLE_KRW = '과세가격원화'
COL_TAXABLE_USD = '과세가격달러'
COL_LAW_CODE = '법령코드'
COL_ISSUED_DOC_NAME = '발급서류명'
COL_NON_TARGET_REASON = '비대상사유'
COL_FREIGHT = '운임'
COL_FREIGHT_CURRENCY = '운임통화단위'
COL_INPUT_FREIGHT = '입력운임'
COL_CALCULATED_FREIGHT_KRW = '계산된운임원화'
COL_INCOTERMS = '인도조건'
COL_TARIFF_EXEMPTION_CODE = '관세감면분납부호'
COL_TARIFF_EXEMPTION_RATE = '관세감면율'

# --- Page Configuration ---
st.set_page_config(
    page_title="TradeGuard - WATI Import 지능형 수입신고 분석",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Utility Functions ---

def safe_numeric_conversion(series):
    """Safely convert a series to numeric, handling commas and NaNs."""
    if pd.api.types.is_numeric_dtype(series):
        return series.fillna(0)
    return pd.to_numeric(
        series.astype(str).str.replace(',', '').fillna('0'),
        errors='coerce'
    ).fillna(0)

def normalize_column_names(df):
    """Strip whitespace from column names and handle duplicates."""
    df.columns = df.columns.str.strip()
    
    # Handle duplicate columns
    cols = pd.Series(df.columns)
    for dup in cols[cols.duplicated()].unique():
        dup_indices = cols[cols == dup].index.tolist()
        for i, idx in enumerate(dup_indices):
            if i > 0:
                cols.iloc[idx] = f"{dup}_{i}"
    df.columns = cols.tolist()
    return df

def map_columns(df):
    """Ensure required columns exist, mapping them if necessary."""
    # 1. Check for exact matches first
    has_rate_type = COL_RATE_TYPE in df.columns
    has_tariff_rate = COL_TARIFF_RATE in df.columns
    
    # 2. Try to find by keywords if missing
    if not has_rate_type:
        candidates = [c for c in df.columns if '세율' in c and '구분' in c]
        if candidates:
            df.rename(columns={candidates[0]: COL_RATE_TYPE}, inplace=True)
            has_rate_type = True
            
    if not has_tariff_rate:
        candidates = [c for c in df.columns if '관세' in c and '율' in c and c != COL_RATE_TYPE]
        if candidates:
             df.rename(columns={candidates[0]: COL_TARIFF_RATE}, inplace=True)
             has_tariff_rate = True
             
    # 운임 컬럼 매핑 시도
    if COL_FREIGHT not in df.columns:
        candidates = [c for c in df.columns if '운임' in c]
        if candidates:
            df.rename(columns={candidates[0]: COL_FREIGHT}, inplace=True)
            
    # 무역거래처국가코드 매핑 시도 (해외공급자 국가코드 등)
    if COL_TRADE_COUNTRY not in df.columns:
        candidates = [c for c in df.columns if '국가코드' in c and '거래처' in c] # 거래처국가코드
        if not candidates:
            candidates = [c for c in df.columns if '해외공급자' in c and '국가' in c] # 해외공급자 국가코드
        if not candidates:
            candidates = [c for c in df.columns if '적출국' in c] # 차선책: 적출국
            
        if candidates:
            df.rename(columns={candidates[0]: COL_TRADE_COUNTRY}, inplace=True)

    # 3. Set defaults if still missing
    if not has_rate_type:
        df[COL_RATE_TYPE] = 'A'
        
    if not has_tariff_rate:
        df[COL_TARIFF_RATE] = 0
        
    if COL_FREIGHT not in df.columns:
        df[COL_FREIGHT] = 0 # Default if not found
        
    return df

def calculate_duty_per_row(df):
    """Calculate '행별관세': (실제관세액 * 금액) / 란결제금액"""
    required = [COL_ACTUAL_DUTY, COL_AMOUNT, COL_LINE_PAYMENT_AMT]
    if all(col in df.columns for col in required):
        for col in required:
            df[col] = safe_numeric_conversion(df[col])
            
        return np.where(
            df[COL_LINE_PAYMENT_AMT] != 0,
            (df[COL_ACTUAL_DUTY] * df[COL_AMOUNT]) / df[COL_LINE_PAYMENT_AMT],
            0
        )
    return 0

def format_date_columns(df):
    """날짜 컬럼을 정수 형식으로 포맷팅 (20250102.0 -> 20250102)"""
    df_display = df.copy()
    date_columns = [COL_ACCEPTANCE_DATE]  # 수리일자
    
    for col in date_columns:
        if col in df_display.columns:
            # 숫자형이면 정수로 변환 후 문자열로
            try:
                df_display[col] = pd.to_numeric(df_display[col], errors='coerce').fillna(0).astype(int).astype(str)
                # 0을 빈 문자열로 변경
                df_display[col] = df_display[col].replace('0', '')
            except:
                pass
    
    return df_display

# --- Main Logic ---

def read_excel_file(uploaded_file, progress_bar=None, status_text=None):
    """Read and preprocess the uploaded Excel file."""
    try:
        if status_text: status_text.text("📂 엑셀 파일 로드 중...")
        if progress_bar: progress_bar.progress(20)
        
        # 파일 확장자 확인 및 로드 방식 결정
        if uploaded_file.name.endswith('.csv'):
             df = pd.read_csv(uploaded_file)
        else:
             df = pd.read_excel(uploaded_file)
        
        if status_text: status_text.text(f"📊 데이터 로드 완료: {len(df):,}행, {len(df.columns)}열")
        if progress_bar: progress_bar.progress(40)
        
        # Normalize columns
        if status_text: status_text.text("🔧 컬럼명 정리 중...")
        df = normalize_column_names(df)
        if progress_bar: progress_bar.progress(60)
        
        # Map columns
        if status_text: status_text.text("🏷️ 컬럼 매핑 중...")
        df = map_columns(df)
        if progress_bar: progress_bar.progress(80)
        
        # Convert types
        if status_text: status_text.text("🔢 데이터 타입 변환 중...")
        if COL_TARIFF_RATE in df.columns:
            df[COL_TARIFF_RATE] = safe_numeric_conversion(df[COL_TARIFF_RATE])
            
        if progress_bar: progress_bar.progress(100)
        if status_text: status_text.text("✅ 데이터 처리 완료!")
        
        return df
    except Exception as e:
        if status_text: status_text.text(f"❌ 오류 발생: {str(e)}")
        st.error(f"파일 읽기 실패: {str(e)}")
        st.error("파일 형식을 확인하거나 다른 파일을 시도해보세요.")
        return None

# --- Existing Analysis Functions ---

def create_eight_percent_refund_analysis(df):
    """8% 환급 검토 분석"""
    try:
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [
            COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
            COL_HS_CODE, COL_RATE_TYPE, COL_RATE_DESC, COL_TARIFF_RATE, 
            COL_EXPORT_COUNTRY, COL_ORIGIN_COUNTRY, COL_FTA_REVIEW, 
            COL_SPEC_1, COL_SPEC_2, COL_SPEC_3, COL_COMP_1, COL_COMP_2, COL_COMP_3,
            COL_ACTUAL_DUTY, COL_PAYMENT_METHOD, COL_CURRENCY, COL_TRADE_NAME, 
            COL_LINE_NO, COL_ROW_NO, COL_QTY_1, COL_UNIT_1, COL_UNIT_PRICE, 
            COL_AMOUNT, COL_LINE_PAYMENT_AMT, COL_ROW_DUTY
        ]
        
        available_cols = [c for c in target_cols if c in df.columns and c not in [COL_ROW_DUTY, COL_FTA_REVIEW]]
        df_work = df[available_cols].copy()
        
        df_work[COL_RATE_TYPE] = df_work[COL_RATE_TYPE].astype(str).str.strip()
        df_work[COL_TARIFF_RATE] = safe_numeric_conversion(df_work[COL_TARIFF_RATE])
        
        df_work[COL_ROW_DUTY] = calculate_duty_per_row(df_work)
        
        if COL_EXPORT_COUNTRY in df_work.columns and COL_ORIGIN_COUNTRY in df_work.columns:
            df_work[COL_FTA_REVIEW] = df_work.apply(
                lambda row: 'FTA사후환급 검토' if (
                    pd.notna(row[COL_EXPORT_COUNTRY]) and 
                    pd.notna(row[COL_ORIGIN_COUNTRY]) and 
                    str(row[COL_EXPORT_COUNTRY]).strip() == str(row[COL_ORIGIN_COUNTRY]).strip() and
                    str(row[COL_EXPORT_COUNTRY]).strip() != ''
                ) else '', axis=1
            )
        else:
            df_work[COL_FTA_REVIEW] = ''
            
        df_filtered = df_work[
            (df_work[COL_RATE_TYPE] == 'A')
        ]
        
        final_cols = [c for c in target_cols if c in df_filtered.columns and c != COL_LINE_PAYMENT_AMT]
        return df_filtered[final_cols].fillna(0)
        
    except Exception as e:
        st.error(f"8% 환급 검토 분석 중 오류 발생: {str(e)}")
        return pd.DataFrame()

def create_zero_percent_risk_analysis(df):
    """0% Risk 분석"""
    try:
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [
            COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
            COL_HS_CODE, COL_RATE_TYPE, COL_TARIFF_RATE, COL_SPEC_1, COL_SPEC_2, 
            COL_COMP_1, COL_ACTUAL_DUTY, COL_TRADE_NAME, COL_LINE_NO, COL_ROW_NO, 
            COL_QTY_1, COL_UNIT_1, COL_UNIT_PRICE, COL_AMOUNT, COL_LINE_PAYMENT_AMT, COL_ROW_DUTY
        ]
        
        df['세율구분_str'] = df[COL_RATE_TYPE].astype(str).str.strip()
        df_zero_risk = df[
            (df[COL_TARIFF_RATE] < 8) & 
            (~df['세율구분_str'].str.match(r'^F.{3}$')) & 
            (~df['세율구분_str'].str.startswith('FR'))
        ].copy()
        
        available_cols = [c for c in target_cols if c in df_zero_risk.columns and c != COL_ROW_DUTY]
        df_zero_risk = df_zero_risk[available_cols].copy()
        
        df_zero_risk[COL_ROW_DUTY] = calculate_duty_per_row(df_zero_risk)
        
        final_cols = [c for c in target_cols if c in df_zero_risk.columns and c != COL_LINE_PAYMENT_AMT]
        return df_zero_risk[final_cols].fillna(0)
        
    except Exception as e:
        st.error(f"0% Risk 분석 중 오류 발생: {str(e)}")
        return pd.DataFrame()

def create_tariff_risk_analysis(df):
    """세율 Risk 분석"""
    try:
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        required_cols = [
            COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
            COL_SPEC_1, COL_SPEC_2, COL_SPEC_3, COL_COMP_1, COL_COMP_2, COL_COMP_3, 
            COL_HS_CODE, COL_RATE_TYPE, COL_RATE_DESC, COL_TAXABLE_USD, COL_ACTUAL_DUTY, 
            COL_PAYMENT_METHOD, COL_AMOUNT, COL_LINE_PAYMENT_AMT, COL_TRADE_NAME
        ]
        
        if COL_SPEC_1 not in df.columns or COL_HS_CODE not in df.columns:
            return pd.DataFrame()
            
        risk_specs = df.groupby(COL_SPEC_1)[COL_HS_CODE].nunique()
        risk_specs = risk_specs[risk_specs > 1]
        
        if len(risk_specs) == 0:
            return pd.DataFrame()
            
        available_cols = [c for c in required_cols if c in df.columns]
        risk_data = df[df[COL_SPEC_1].isin(risk_specs.index)][available_cols].copy()
        
        risk_data[COL_ROW_DUTY] = calculate_duty_per_row(risk_data)
        
        risk_data = risk_data.sort_values([COL_SPEC_1, COL_HS_CODE]).fillna('')
        
        # 관세실행세율 추가
        display_cols = [COL_SPEC_1, COL_HS_CODE, COL_TARIFF_RATE, COL_TAX_CLASSIFICATION, COL_TRADE_NAME]
        final_cols = [c for c in available_cols if c != COL_LINE_PAYMENT_AMT]
        
        # display_cols에 있는 컬럼 중 available_cols에 없는 것 추가
        for col in display_cols:
            if col in df.columns and col not in final_cols:
                final_cols.insert(min(2, len(final_cols)), col)
        
        if COL_ROW_DUTY not in final_cols:
            final_cols.append(COL_ROW_DUTY)
            
        return risk_data[final_cols]
        
    except Exception as e:
        st.error(f"세율 Risk 분석 중 오류 발생: {e}")
        return pd.DataFrame()

def create_price_risk_analysis(df):
    """단가 Risk 분석 (Z-Score 기반)"""
    try:
        if COL_UNIT_PRICE not in df.columns or COL_SPEC_1 not in df.columns:
            return pd.DataFrame()
            
        # 분석에 필요한 컬럼 정의
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_HS_CODE, COL_TRADE_NAME, COL_SPEC_1, COL_UNIT_PRICE, COL_CURRENCY, COL_AMOUNT, COL_QTY_1]
        available_cols = [c for c in target_cols if c in df.columns]
        
        df_work = df.copy()
        df_work[COL_UNIT_PRICE] = safe_numeric_conversion(df_work[COL_UNIT_PRICE])
        # 단가가 0보다 큰 건만 분석
        df_work = df_work[df_work[COL_UNIT_PRICE] > 0]
        
        if len(df_work) == 0:
            return pd.DataFrame()

        # 규격1별 통계 산출 (평균, 표준편차)
        # 데이터 개수(count)가 적으면(예: 3개 미만) 통계적 유의성이 낮으므로 Z-Score 계산에서 제외하거나 주의 필요
        # 여기서는 최소 3건 이상인 규격만 분석 대상으로 삼음
        stats = df_work.groupby(COL_SPEC_1)[COL_UNIT_PRICE].agg(['mean', 'std', 'count']).reset_index()
        stats = stats[stats['count'] >= 3] 

        if len(stats) == 0:
            return pd.DataFrame()

        # 원본 데이터에 통계 매핑
        df_merged = pd.merge(df_work, stats, on=COL_SPEC_1, how='inner')

        # Z-Score 계산: (단가 - 평균) / 표준편차
        # 표준편차가 0인 경우(모든 단가가 동일)는 Z-Score 0으로 처리
        df_merged['z_score'] = np.where(
            df_merged['std'] > 0,
            (df_merged[COL_UNIT_PRICE] - df_merged['mean']) / df_merged['std'],
            0
        )

        # 이상치 필터링 (Z-Score 절대값이 1.96 이상인 경우 - 95% 신뢰구간 밖)
        # 1.96은 통계적으로 유의미한 이상치 기준 중 하나 (약 상위/하위 2.5%)
        threshold = 1.96 
        outliers = df_merged[abs(df_merged['z_score']) > threshold].copy()
        
        if len(outliers) == 0:
            return pd.DataFrame()

        # 보기 좋게 반올림
        outliers['Z-Score'] = outliers['z_score'].round(2)
        outliers['평균단가'] = outliers['mean'].round(2)
        outliers['표준편차'] = outliers['std'].round(2)
        
        # 정렬: Z-Score 절대값이 높은 순서대로 (가장 이상한 것부터)
        outliers = outliers.sort_values(by='Z-Score', key=abs, ascending=False)
        
        final_cols = available_cols + ['Z-Score', '평균단가', '표준편차']
        return outliers[final_cols]
        
    except Exception as e:
        st.error(f"단가 Risk 분석 중 오류 발생: {str(e)}")
        return pd.DataFrame()

def create_domestic_tax_code_analysis(df):
    """내국세구분 분석"""
    try:
        if COL_HS_CODE not in df.columns:
            return pd.DataFrame()
            
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [
            COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
            COL_HS_CODE, COL_RATE_TYPE, COL_TARIFF_RATE, COL_INTERNAL_TAX_CODE, 
            COL_SPEC_1, COL_SPEC_2, COL_SPEC_3, COL_COMP_1, COL_COMP_2, COL_COMP_3, 
            COL_ACTUAL_DUTY, COL_TRADE_NAME, COL_LINE_NO, COL_ROW_NO, COL_QTY_1, 
            COL_UNIT_1, COL_UNIT_PRICE, COL_AMOUNT, COL_LINE_PAYMENT_AMT, COL_ROW_DUTY
        ]
        
        available_cols = [c for c in target_cols if c in df.columns and c not in [COL_ROW_DUTY, COL_INTERNAL_TAX_CODE]]
        df_work = df[available_cols].copy()
        
        if COL_INTERNAL_TAX_CODE not in df.columns:
            df_work[COL_INTERNAL_TAX_CODE] = ''
        else:
            df_work[COL_INTERNAL_TAX_CODE] = df[COL_INTERNAL_TAX_CODE].fillna('').astype(str).str.strip()
            
        df_work[COL_HS_CODE] = df_work[COL_HS_CODE].astype(str).str.strip()
        
        df_filtered = df_work[
            (df_work[COL_HS_CODE].str.len() == 10) &
            (df_work[COL_HS_CODE].str.startswith('22')) &
            ((df_work[COL_INTERNAL_TAX_CODE] == '') | (df_work[COL_INTERNAL_TAX_CODE].isna()))
        ].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
            
        df_filtered[COL_ROW_DUTY] = calculate_duty_per_row(df_filtered)
        
        final_cols = [c for c in target_cols if c in df_filtered.columns and c != COL_LINE_PAYMENT_AMT]
        return df_filtered[final_cols].fillna(0).sort_values(COL_IMPORT_DEC_NO)
        
    except Exception as e:
        st.error(f"내국세구분 분석 중 오류 발생: {str(e)}")
        return pd.DataFrame()

def create_import_requirement_risk_analysis(df):
    """수입요건 Risk 분석: 동일 규격1 내에서 신고별 법령 세트가 다른 경우 탐지 (개선 버전)"""
    try:
        if COL_SPEC_1 not in df.columns or COL_IMPORT_DEC_NO not in df.columns:
            return pd.DataFrame()
        
        # 분석 대상 컬럼 확인
        req_cols = [COL_LAW_CODE, COL_ISSUED_DOC_NAME, COL_NON_TARGET_REASON]
        available_cols = [col for col in req_cols if col in df.columns]
        
        if not available_cols:
            return pd.DataFrame()
        
        df_work = df.copy()
        risk_declarations = []
        
        # 규격1별로 그룹화
        for spec in df_work[COL_SPEC_1].dropna().unique():
            spec_data = df_work[df_work[COL_SPEC_1] == spec]
            
            # 신고번호별로 그룹화 (신고 단위로 비교)
            declaration_groups = spec_data.groupby(COL_IMPORT_DEC_NO)
            
            # 최소 2개 신고가 있어야 비교 가능
            if len(declaration_groups) < 2:
                continue
            
            # 각 신고별로 법령 세트 생성
            declaration_sets = {}
            for decl_no, decl_data in declaration_groups:
                # 각 컬럼별로 고유값을 tuple로 만들어서 세트화
                law_set = set()
                for col in available_cols:
                    vals = decl_data[col].dropna()
                    vals = vals[vals.astype(str).str.strip() != '']
                    if len(vals) > 0:
                        # 여러 행이 있어도 고유값만 추출
                        law_set.update(vals.unique())
                
                # 세트를 frozenset으로 변환 (비교 가능하도록)
                if law_set:  # 비어있지 않은 경우만
                    declaration_sets[decl_no] = frozenset(law_set)
            
            # 신고별 세트가 2개 이상이고, 서로 다른 세트가 존재하는지 확인
            if len(declaration_sets) >= 2:
                unique_sets = set(declaration_sets.values())
                
                # 서로 다른 세트가 2개 이상이면 불일치
                if len(unique_sets) > 1:
                    # 이 규격1의 모든 신고를 위험으로 표시
                    risk_declarations.extend(spec_data[COL_IMPORT_DEC_NO].unique())
        
        if not risk_declarations:
            return pd.DataFrame()
        
        # 위험 신고들의 상세 내역 반환
        df_result = df_work[df_work[COL_IMPORT_DEC_NO].isin(risk_declarations)].copy()
        df_result = df_result.sort_values(by=[COL_SPEC_1, COL_IMPORT_DEC_NO])
        
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        display_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                        COL_SPEC_1, COL_HS_CODE] + available_cols + [COL_TRADE_NAME, COL_ORIGIN_COUNTRY]
        final_cols = [col for col in display_cols if col in df_result.columns]
        
        return df_result[final_cols]
        
    except Exception as e:
        st.error(f"수입요건 Risk 분석 오류: {str(e)}")
        return pd.DataFrame()

# --- New Analysis Functions (Requested 12-19) ---

def create_f_rate_analysis(df):
    """12. F세율 (F세율이 적용되는 세번 선별)"""
    try:
        if COL_RATE_TYPE not in df.columns:
            return pd.DataFrame()
            
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_HS_CODE, COL_RATE_TYPE, COL_RATE_DESC, 
                       COL_TARIFF_RATE, COL_TRADE_NAME, COL_SPEC_1, COL_ORIGIN_COUNTRY, COL_AMOUNT]
        available_cols = [c for c in target_cols if c in df.columns]
        
        # 세율구분이 정확히 'F' (한 글자)인 건만 필터링
        df_filtered = df[df[COL_RATE_TYPE].astype(str).str.strip() == 'F'].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
        
        # 행별관세 계산
        df_filtered[COL_ROW_DUTY] = calculate_duty_per_row(df_filtered)
        
        # 최종 컬럼에 행별관세 추가
        final_cols = available_cols + [COL_ROW_DUTY]
        
        return df_filtered[final_cols].fillna('')
    except Exception as e:
        st.error(f"F세율 분석 중 오류: {str(e)}")
        return pd.DataFrame()

def create_fta_opportunity_analysis(df):
    """13. FTA 기회 발굴 (적출국=원산지, A세율 적용 건)"""
    try:
        required_cols = [COL_RATE_TYPE, COL_EXPORT_COUNTRY, COL_ORIGIN_COUNTRY, COL_TARIFF_RATE]
        if not all(c in df.columns for c in required_cols):
            return pd.DataFrame()
            
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_HS_CODE, COL_RATE_TYPE, COL_TARIFF_RATE, COL_EXPORT_COUNTRY, 
                       COL_ORIGIN_COUNTRY, COL_TAXABLE_USD, COL_TRADE_NAME, COL_AMOUNT, COL_ACTUAL_DUTY]
        available_cols = [c for c in target_cols if c in df.columns]
        
        # 조건: 세율구분 'A' & 적출국 == 원산지 & 관세율 > 0
        df_work = df.copy()
        df_work[COL_TARIFF_RATE] = safe_numeric_conversion(df_work[COL_TARIFF_RATE])
        
        df_filtered = df_work[
            (df_work[COL_RATE_TYPE].astype(str).str.strip() == 'A') &
            (df_work[COL_EXPORT_COUNTRY].fillna('').astype(str).str.strip() == df_work[COL_ORIGIN_COUNTRY].fillna('').astype(str).str.strip()) &
            (df_work[COL_EXPORT_COUNTRY].fillna('').astype(str).str.strip() != '') &
            (df_work[COL_TARIFF_RATE] > 0)
        ].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
            
        # 과세가격 기준 내림차순 정렬 (우선순위)
        if COL_TAXABLE_USD in df_filtered.columns:
             df_filtered[COL_TAXABLE_USD] = safe_numeric_conversion(df_filtered[COL_TAXABLE_USD])
             df_filtered = df_filtered.sort_values(by=COL_TAXABLE_USD, ascending=False)
             
        return df_filtered[available_cols]
    except Exception as e:
        st.error(f"FTA 기회 발굴 분석 중 오류: {str(e)}")
        return pd.DataFrame()

def create_low_price_analysis(df, threshold=10):
    """14. 과세가격 (단가가 낮은 신고건 선별 - 저가신고 우려)"""
    try:
        if COL_UNIT_PRICE not in df.columns:
            return pd.DataFrame()
            
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_HS_CODE, COL_TRADE_NAME, COL_SPEC_1, COL_UNIT_PRICE, 
                       COL_CURRENCY, COL_AMOUNT, COL_PAYMENT_METHOD]
        available_cols = [c for c in target_cols if c in df.columns]
        
        df_work = df.copy()
        df_work[COL_UNIT_PRICE] = safe_numeric_conversion(df_work[COL_UNIT_PRICE])
        
        # 단가가 threshold(기본 10) 이하인 건
        df_filtered = df_work[df_work[COL_UNIT_PRICE] <= threshold].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
            
        return df_filtered[available_cols].sort_values(by=COL_UNIT_PRICE)
    except Exception as e:
        st.error(f"저가신고 분석 중 오류: {str(e)}")
        return pd.DataFrame()

def create_currency_consistency_analysis(df):
    """15. 통화단위 (무역거래처별 통화단위 일관성 + 이상치점수)"""
    try:
        if COL_TRADE_COMPANY not in df.columns or COL_CURRENCY not in df.columns:
            return pd.DataFrame()
            
        # 거래처별 사용 통화 집계
        grouped = df.groupby(COL_TRADE_COMPANY)[COL_CURRENCY].unique().reset_index()
        grouped['통화개수'] = grouped[COL_CURRENCY].apply(len)
        
        # 통화가 2개 이상인 거래처 필터링
        inconsistent_companies = grouped[grouped['통화개수'] > 1][COL_TRADE_COMPANY].tolist()
        
        if not inconsistent_companies:
            return pd.DataFrame()
            
        df_filtered = df[df[COL_TRADE_COMPANY].isin(inconsistent_companies)].copy()
        
        # 이상치점수 계산 (거래처-통화 조합별 빈도 기반)
        if COL_TRADE_COUNTRY in df.columns:
            # 국가-통화 조합별 빈도 계산
            country_currency_counts = df.groupby([COL_TRADE_COUNTRY, COL_CURRENCY]).size().reset_index(name='count')
            country_totals = df.groupby(COL_TRADE_COUNTRY).size().reset_index(name='total')
            merged = pd.merge(country_currency_counts, country_totals, on=COL_TRADE_COUNTRY)
            merged['ratio'] = merged['count'] / merged['total']
            merged['이상치점수'] = ((1 - merged['ratio']) * 100).round(1)
            
            # 결과에 이상치점수 추가
            df_filtered = pd.merge(df_filtered, 
                                  merged[[COL_TRADE_COUNTRY, COL_CURRENCY, '이상치점수']], 
                                  on=[COL_TRADE_COUNTRY, COL_CURRENCY], 
                                  how='left')
        
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_CURRENCY, '이상치점수', COL_AMOUNT]
        available_cols = [c for c in target_cols if c in df_filtered.columns]
        
        df_filtered = df_filtered.sort_values(by=[COL_TRADE_COMPANY, COL_CURRENCY])
        
        return df_filtered[available_cols]
    except Exception as e:
        st.error(f"통화단위 일관성 분석 중 오류: {str(e)}")
        return pd.DataFrame()

# 국가별 통화단위 불일치 분석 - 제거됨 (통화단위 불일치에 통합)
# def create_country_currency_consistency_analysis(df):
    """국가별 통화단위 불일치 분석 (빈도 기반 이상치 탐지)"""
    try:
        if COL_TRADE_COUNTRY not in df.columns or COL_CURRENCY not in df.columns:
            return pd.DataFrame()
        
        df_work = df.dropna(subset=[COL_TRADE_COUNTRY, COL_CURRENCY]).copy()
        
        # 국가별 전체 신고 건수 계산
        country_counts = df_work[COL_TRADE_COUNTRY].value_counts().reset_index()
        country_counts.columns = [COL_TRADE_COUNTRY, 'total_count']
        
        # 국가별 통화별 신고 건수 계산
        currency_counts = df_work.groupby([COL_TRADE_COUNTRY, COL_CURRENCY]).size().reset_index(name='count')
        
        # 병합
        merged = pd.merge(currency_counts, country_counts, on=COL_TRADE_COUNTRY)
        
        # 비율 계산
        merged['ratio'] = merged['count'] / merged['total_count']
        
        # 이상치 점수 (Anomaly Score) 계산: (1 - 비율) * 100
        # 희귀할수록 점수가 100에 가까워짐
        merged['anomaly_score'] = (1 - merged['ratio']) * 100
        
        # 통화가 1개뿐인 국가는 제외 (일관성 문제 없음)
        multi_currency_countries = merged[merged['total_count'] > merged['count']][COL_TRADE_COUNTRY].unique()
        merged = merged[merged[COL_TRADE_COUNTRY].isin(multi_currency_countries)]
        
        # 이상치 필터링 (비율 10% 미만인 경우를 '튀는 것'으로 간주)
        outliers = merged[merged['ratio'] < 0.1].copy()
        
        if len(outliers) == 0:
            return pd.DataFrame()
        
        # 원본 데이터와 병합하여 상세 정보 표시
        # outlier에 해당하는 국가-통화 조합만 추출
        target_keys = outliers[[COL_TRADE_COUNTRY, COL_CURRENCY]].drop_duplicates()
        
        df_result = pd.merge(df_work, target_keys, on=[COL_TRADE_COUNTRY, COL_CURRENCY], how='inner')
        
        # 점수 정보 추가
        df_result = pd.merge(df_result, outliers[[COL_TRADE_COUNTRY, COL_CURRENCY, 'ratio', 'anomaly_score']], on=[COL_TRADE_COUNTRY, COL_CURRENCY], how='left')
        
        # 보기 좋게 포맷팅
        df_result['사용비율'] = (df_result['ratio'] * 100).round(1).astype(str) + '%'
        df_result['이상치점수'] = df_result['anomaly_score'].round(1)
        
        target_cols = [COL_TRADE_COUNTRY, COL_CURRENCY, '사용비율', '이상치점수', COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_AMOUNT]
        available_cols = [c for c in target_cols if c in df_result.columns]
        
        return df_result[available_cols].sort_values(by='이상치점수', ascending=False)
        
    except Exception as e:
        st.error(f"국가별 통화단위 분석 중 오류: {str(e)}")
        return pd.DataFrame()

# def create_trade_type_consistency_analysis(df):
#     """특수거래 구분 분석 - 사용자 요청으로 제거됨"""
#     return pd.DataFrame()

def create_free_charge_freight_analysis(df):
    """17. 무상 운임 누락 (인도조건 EXW/FOB이면서 입력운임이 없는 경우)"""
    try:
        # 인도조건 컬럼 필수
        if COL_INCOTERMS not in df.columns:
            return pd.DataFrame()
            
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_PAYMENT_METHOD, COL_INCOTERMS]
        
        # 운임 관련 컬럼 추가
        freight_cols = [COL_FREIGHT, COL_FREIGHT_CURRENCY, COL_INPUT_FREIGHT, COL_CALCULATED_FREIGHT_KRW]
        for col in freight_cols:
            if col in df.columns:
                target_cols.append(col)
        
        target_cols.extend([COL_AMOUNT, COL_TRADE_NAME])
        
        available_cols = [c for c in target_cols if c in df.columns]
        
        df_work = df.copy()
        
        # 인도조건이 EXW 또는 FOB이면서 입력운임이 없는 경우
        incoterms_condition = df_work[COL_INCOTERMS].astype(str).str.strip().isin(['EXW', 'FOB'])
        
        # 입력운임 컬럼이 있는 경우
        if COL_INPUT_FREIGHT in df_work.columns:
            # 입력운임이 비어있거나 0인 경우
            input_freight_missing = (
                df_work[COL_INPUT_FREIGHT].isna() | 
                (df_work[COL_INPUT_FREIGHT].astype(str).str.strip() == '') |
                (safe_numeric_conversion(df_work[COL_INPUT_FREIGHT]) == 0)
            )
        else:
            # 입력운임 컬럼 자체가 없으면 모두 누락으로 간주
            input_freight_missing = True
        
        df_filtered = df_work[incoterms_condition & input_freight_missing].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
            
        return df_filtered[available_cols]
    except Exception as e:
        st.error(f"무상 운임 누락 분석 중 오류: {str(e)}")
        return pd.DataFrame()

def create_usage_rate_analysis(df):
    """19. 용도세율 (HSK 코드 기반 용도세율 적용 품목 선별)"""
    try:
        if COL_HS_CODE not in df.columns:
            return pd.DataFrame()
        
        # CSV 파일에서 HSK 코드 목록 읽기
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            csv_path = os.path.join(current_dir, 'usage_rate_hsk.csv')
            
            if os.path.exists(csv_path):
                hsk_df = pd.read_csv(csv_path, encoding='utf-8')
            else:
                st.warning("용도세율 HSK 파일(usage_rate_hsk.csv)을 찾을 수 없습니다.")
                return pd.DataFrame()
            
            # HSK 컬럼의 점과 하이픈 제거하여 10자리 숫자만 추출
            hsk_df['HSK_10'] = hsk_df['HSK'].astype(str).str.replace('.', '').str.replace('-', '')
            hsk_dict = dict(zip(hsk_df['HSK_10'], zip(hsk_df['용도'], hsk_df['출처'])))
            
        except Exception as e:
            st.error(f"HSK CSV 파일 로드 중 오류: {str(e)}")
            return pd.DataFrame()
        
        # 데이터의 세번부호 10자리 추출
        df_work = df.copy()
        df_work['세번부호_10'] = df_work[COL_HS_CODE].astype(str).str.replace('.', '').str.replace('-', '').str[:10]
        
        # HSK 목록과 매칭
        df_work['매칭여부'] = df_work['세번부호_10'].isin(hsk_dict.keys())
        df_filtered = df_work[df_work['매칭여부']].copy()
        
        if len(df_filtered) == 0:
            return pd.DataFrame()
        
        # 용도 및 출처 정보 추가
        df_filtered['용도'] = df_filtered['세번부호_10'].map(lambda x: hsk_dict.get(x, ('', ''))[0])
        df_filtered['출처'] = df_filtered['세번부호_10'].map(lambda x: hsk_dict.get(x, ('', ''))[1])
        
        # 공통 최우선 컬럼 + 특정 분석 컬럼
        target_cols = [COL_IMPORT_DEC_NO, COL_ACCEPTANCE_DATE, COL_BL_NO, COL_TRADE_COMPANY, COL_TRADE_COUNTRY,
                       COL_HS_CODE, '용도', '출처', COL_RATE_TYPE, COL_RATE_DESC, COL_TARIFF_RATE, 
                       COL_TRADE_NAME, COL_SPEC_1, COL_AMOUNT]
        available_cols = [c for c in target_cols if c in df_filtered.columns]
        
        return df_filtered[available_cols]
        
    except Exception as e:
        st.error(f"용도세율 분석 중 오류: {str(e)}")
        return pd.DataFrame()

def create_summary_analysis(df_original):
    """Summary 분석"""
    try:
        summary_data = {}
        
        if COL_IMPORT_DEC_NO in df_original.columns:
            total_declarations = df_original[COL_IMPORT_DEC_NO].nunique()
        else:
            total_declarations = len(df_original)
        summary_data['전체 신고 건수'] = total_declarations
        
        if COL_TRADE_TYPE in df_original.columns and COL_IMPORT_DEC_NO in df_original.columns:
            trade_type_analysis = pd.pivot_table(df_original, 
                index=[COL_TRADE_TYPE], values=COL_IMPORT_DEC_NO, aggfunc='nunique',
                margins=True, margins_name='총계'
            ).reset_index()
            summary_data['거래구분별'] = trade_type_analysis
            
        if COL_RATE_TYPE in df_original.columns and COL_IMPORT_DEC_NO in df_original.columns:
            rate_type_analysis = pd.pivot_table(df_original,
                index=COL_RATE_TYPE, values=COL_IMPORT_DEC_NO, aggfunc='nunique'
            ).reset_index()
            total_row = {COL_RATE_TYPE: '총계', COL_IMPORT_DEC_NO: rate_type_analysis[COL_IMPORT_DEC_NO].sum()}
            rate_type_analysis = pd.concat([rate_type_analysis, pd.DataFrame([total_row])], ignore_index=True)
            summary_data['세율구분별'] = rate_type_analysis
            
        # 기존 Risk Counts 계산
        zero_risk_count = 0
        if all(col in df_original.columns for col in [COL_TARIFF_RATE, COL_RATE_TYPE, COL_IMPORT_DEC_NO]):
            df_original['세율구분_str'] = df_original[COL_RATE_TYPE].astype(str).str.strip()
            zero_risk_df = df_original[
                (df_original[COL_TARIFF_RATE] < 8) & 
                (~df_original['세율구분_str'].str.match(r'^F.{3}$')) &
                (~df_original['세율구분_str'].str.startswith('FR'))
            ]
            zero_risk_count = zero_risk_df[COL_IMPORT_DEC_NO].nunique()
            
        eight_percent_count = 0
        if all(col in df_original.columns for col in [COL_TARIFF_RATE, COL_RATE_TYPE, COL_IMPORT_DEC_NO]):
            eight_percent_df = df_original[
                (df_original[COL_RATE_TYPE] == 'A') & 
                (df_original[COL_TARIFF_RATE] >= 8)
            ]
            eight_percent_count = eight_percent_df[COL_IMPORT_DEC_NO].nunique()
            
        tariff_risk_count = 0
        if COL_SPEC_1 in df_original.columns and COL_HS_CODE in df_original.columns:
            risk_specs = df_original.groupby(COL_SPEC_1)[COL_HS_CODE].nunique()
            risk_specs = risk_specs[risk_specs > 1]
            if len(risk_specs) > 0:
                tariff_risk_df = df_original[df_original[COL_SPEC_1].isin(risk_specs.index)]
                tariff_risk_count = tariff_risk_df[COL_IMPORT_DEC_NO].nunique()
        
        domestic_tax_count = 0
        if COL_HS_CODE in df_original.columns:
            df_tax = df_original.copy()
            if COL_INTERNAL_TAX_CODE not in df_tax.columns:
                df_tax[COL_INTERNAL_TAX_CODE] = ''
            else:
                df_tax[COL_INTERNAL_TAX_CODE] = df_tax[COL_INTERNAL_TAX_CODE].fillna('').astype(str).str.strip()
            
            df_tax[COL_HS_CODE] = df_tax[COL_HS_CODE].astype(str).str.strip()
            domestic_tax_df = df_tax[
                (df_tax[COL_HS_CODE].str.len() == 10) &
                (df_tax[COL_HS_CODE].str.startswith('22')) &
                ((df_tax[COL_INTERNAL_TAX_CODE] == '') | (df_tax[COL_INTERNAL_TAX_CODE].isna()))
            ]
            domestic_tax_count = domestic_tax_df[COL_IMPORT_DEC_NO].nunique()
        
        price_risk_count = 0
        if COL_SPEC_1 in df_original.columns and COL_UNIT_PRICE in df_original.columns:
            # Z-Score 기반 카운팅
            df_price = df_original.copy()
            df_price[COL_UNIT_PRICE] = safe_numeric_conversion(df_price[COL_UNIT_PRICE])
            df_price = df_price[df_price[COL_UNIT_PRICE] > 0]
            
            stats = df_price.groupby(COL_SPEC_1)[COL_UNIT_PRICE].agg(['mean', 'std', 'count']).reset_index()
            stats = stats[stats['count'] >= 3]
            
            if not stats.empty:
                df_merged = pd.merge(df_price, stats, on=COL_SPEC_1, how='inner')
                df_merged['z_score'] = np.where(
                    df_merged['std'] > 0,
                    (df_merged[COL_UNIT_PRICE] - df_merged['mean']) / df_merged['std'],
                    0
                )
                price_risk_df = df_merged[abs(df_merged['z_score']) > 1.96]
                price_risk_count = price_risk_df[COL_IMPORT_DEC_NO].nunique()
        
        import_req_risk_count = 0
        if COL_SPEC_1 in df_original.columns:
            req_cols = [COL_LAW_CODE, COL_ISSUED_DOC_NAME, COL_NON_TARGET_REASON]
            available_cols = [col for col in req_cols if col in df_original.columns]
            
            if available_cols:
                risk_specs = []
                for spec in df_original[COL_SPEC_1].dropna().unique():
                    spec_data = df_original[df_original[COL_SPEC_1] == spec]
                    if len(spec_data) >= 2:
                        for col in available_cols:
                            unique_vals = spec_data[col].dropna()
                            unique_vals = unique_vals[unique_vals.astype(str).str.strip() != '']
                            if len(unique_vals.unique()) > 1:
                                risk_specs.append(spec)
                                break
                
                if risk_specs:
                    import_req_df = df_original[df_original[COL_SPEC_1].isin(risk_specs)]
                    import_req_risk_count = import_req_df[COL_IMPORT_DEC_NO].nunique()

        # New Risk Counts
        f_rate_count = len(create_f_rate_analysis(df_original))
        fta_opp_count = len(create_fta_opportunity_analysis(df_original))
        low_price_count = len(create_low_price_analysis(df_original))
        currency_inc_count = len(create_currency_consistency_analysis(df_original))
        country_curr_inc_count = len(create_country_currency_consistency_analysis(df_original)) # New
        trade_type_count = len(create_trade_type_consistency_analysis(df_original))
        free_freight_count = len(create_free_charge_freight_analysis(df_original))
        usage_rate_count = len(create_usage_rate_analysis(df_original))

        risk_analysis = pd.DataFrame({
            'Risk 유형': [
                '0% 세율 위험', '8% 환급 검토', '세율 위험(HS 불일치)', '내국세 누락', '단가 위험(Z-Score)', '수입요건 불일치',
                'F세율 적용', 'FTA 기회 발굴', '저가신고 의심', '통화단위 불일치(거래처)', '국가별 통화단위 불일치', '특수거래 구분', '무상운임 누락', '용도세율 적용'
            ],
            '신고건수': [
                zero_risk_count, eight_percent_count, tariff_risk_count, domestic_tax_count, price_risk_count, import_req_risk_count,
                f_rate_count, fta_opp_count, low_price_count, currency_inc_count, country_curr_inc_count, trade_type_count, free_freight_count, usage_rate_count
            ],
            '비율(%)': [
                zero_risk_count/total_declarations*100 if total_declarations > 0 else 0,
                eight_percent_count/total_declarations*100 if total_declarations > 0 else 0,
                tariff_risk_count/total_declarations*100 if total_declarations > 0 else 0,
                domestic_tax_count/total_declarations*100 if total_declarations > 0 else 0,
                price_risk_count/total_declarations*100 if total_declarations > 0 else 0,
                import_req_risk_count/total_declarations*100 if total_declarations > 0 else 0,
                f_rate_count/total_declarations*100 if total_declarations > 0 else 0,
                fta_opp_count/total_declarations*100 if total_declarations > 0 else 0,
                low_price_count/total_declarations*100 if total_declarations > 0 else 0,
                currency_inc_count/total_declarations*100 if total_declarations > 0 else 0,
                country_curr_inc_count/total_declarations*100 if total_declarations > 0 else 0,
                trade_type_count/total_declarations*100 if total_declarations > 0 else 0,
                free_freight_count/total_declarations*100 if total_declarations > 0 else 0,
                usage_rate_count/total_declarations*100 if total_declarations > 0 else 0
            ]
        })
        summary_data['Risk분석'] = risk_analysis
        
        # 월별 추이 분석
        if COL_ACCEPTANCE_DATE in df_original.columns and COL_IMPORT_DEC_NO in df_original.columns:
            try:
                df_monthly = df_original.copy()
                df_monthly[COL_ACCEPTANCE_DATE] = pd.to_numeric(
                    df_monthly[COL_ACCEPTANCE_DATE], 
                    errors='coerce'
                ).fillna(0).astype(int).astype(str)
                df_monthly = df_monthly[df_monthly[COL_ACCEPTANCE_DATE].str.len() == 8]
                df_monthly[COL_ACCEPTANCE_DATE] = pd.to_datetime(
                    df_monthly[COL_ACCEPTANCE_DATE], 
                    format='%Y%m%d', 
                    errors='coerce'
                )
                df_monthly = df_monthly[df_monthly[COL_ACCEPTANCE_DATE].notna()]
                
                if len(df_monthly) > 0:
                    df_monthly['수리월'] = df_monthly[COL_ACCEPTANCE_DATE].dt.strftime('%Y-%m')
                    monthly_trend = df_monthly.groupby('수리월')[COL_IMPORT_DEC_NO].nunique().reset_index()
                    monthly_trend.columns = ['수리월', '신고건수']
                    monthly_trend = monthly_trend.sort_values('수리월')
                    summary_data['월별추이'] = monthly_trend
            except Exception:
                pass
            
        return summary_data
        
    except Exception as e:
        st.error(f"Summary 분석 중 오류 발생: {str(e)}")
        return {}

def create_verification_methods_excel_sheet(writer):
    """검증방법 시트 생성 (엑셀용)"""
    try:
        worksheet = writer.book.add_worksheet('검증방법')
        workbook = writer.book
        
        title_format = workbook.add_format({
            'font_name': 'Arial', 'font_size': 14, 'bold': True,
            'align': 'center', 'valign': 'vcenter', 'bg_color': '#4472C4',
            'font_color': 'white', 'border': 1
        })
        
        subtitle_format = workbook.add_format({
            'font_name': 'Arial', 'font_size': 12, 'bold': True,
            'align': 'left', 'valign': 'vcenter', 'bg_color': '#D9E1F2', 'border': 1
        })
        
        content_format = workbook.add_format({
            'font_name': 'Arial', 'font_size': 10, 'align': 'left',
            'valign': 'top', 'border': 1, 'text_wrap': True
        })
        
        highlight_format = workbook.add_format({
            'font_name': 'Arial', 'font_size': 10, 'align': 'left',
            'valign': 'top', 'border': 1, 'text_wrap': True, 'bg_color': '#FFFF00'
        })
        
        worksheet.set_column(0, 0, 25)
        worksheet.set_column(1, 1, 60)
        worksheet.set_column(2, 2, 40)
        
        current_row = 0
        worksheet.merge_range(current_row, 0, current_row, 2, '수입신고 분석 검증방법', title_format)
        worksheet.set_row(current_row, 30)
        current_row += 2
        
        sections = [
            ('1. 8% 환급 검토', '• 필터링 조건: 세율구분 = "A" AND 관세실행세율 ≥ 8%\n• 목적: 8% 환급 검토가 필요한 수입신고 건들 식별', '• 세율구분 "A"는 일반적으로 가장 관세율이 높은 구분'),
            ('2. 0% Risk', '• 필터링 조건: 관세실행세율 < 8% AND 세율구분 ≠ F*** AND 세율구분 ≠ FR*', '• 관세율이 낮은데도 특별한 세율구분이 아닌 경우 주의 필요'),
            ('3. 세율 Risk', '• 분석 방법: 규격1 기준으로 그룹화하여 세번부호의 고유값 개수 확인', '• 동일 상품인데 다른 세번부호가 적용되면 관세율 차이 발생'),
            ('4. 단가 Risk (Z-Score)', '• 분석 방법: 규격1별 단가의 표준점수(Z-Score) 산출\n• 기준: |Z-Score| > 1.96 (신뢰구간 95% 밖)', '• 통계적으로 유의미한 단가 이상치 탐지'),
            ('5. 내국세구분', '• 필터링 조건: 세번부호 10자리 AND 22로 시작 AND 내국세부호 없음', '• 주류 수입 시 내국세부호 누락은 세금 신고 오류'),
            ('6. F세율 적용', '• 필터링 조건: 세율구분이 "F"로 시작하는 건', '• FTA 등 협정세율 적용 적정성 확인'),
            ('7. FTA 기회 발굴', '• 필터링 조건: A세율 적용 & 적출국=원산지 & 관세율 > 0', '• FTA 미적용 건 중 적용 가능성 있는 건 발굴'),
            ('8. 저가신고 의심', '• 필터링 조건: 단가 $10 이하', '• 저가 신고로 인한 관세 탈루 리스크 점검'),
            ('9. 통화단위 불일치', '• 거래처별: 동일 거래처 내 통화단위 2개 이상 존재 여부 확인\n• 국가별: 동일 국가 내 통화단위 사용 빈도 분석 (하위 10% 미만 이상치 탐지)', '• 신고 오류 가능성 점검'),
            ('10. 특수거래 구분', '• 필터링 조건: 일반수입(11) 외 거래구분', '• 재수출, 감면 등 특수 거래의 사후관리 필요성 점검'),
            ('11. 무상운임 누락', '• 필터링 조건: 결제방법 "GN"(무상) & 운임 0원', '• 무상 수입 시 운임 누락 여부 점검'),
            ('12. 용도세율 적용', '• 필터링 조건: 세율구분 "C"', '• 용도세율 적용 물품의 사후관리 이행 여부 점검'),
            ('13. 원본데이터', '• 분석에 사용된 원본 엑셀 파일의 모든 데이터', '• 원본 데이터와 분석 결과 비교 검토 가능')
        ]
        
        for title, content, note in sections:
            worksheet.write(current_row, 0, title, subtitle_format)
            worksheet.write(current_row, 1, content, content_format)
            worksheet.write(current_row, 2, note, highlight_format)
            worksheet.set_row(current_row, 80)
            current_row += 1
            
        return True
    except Exception as e:
        print(f"검증방법 시트 생성 중 오류 발생: {str(e)}")
        return False

def create_excel_file(df_original, results, summary_data):
    """Excel 파일 생성 (모든 결과 포함)"""
    try:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            workbook = writer.book
            header_format = workbook.add_format({'bold': True, 'bg_color': '#D9E1F2', 'border': 1, 'align': 'center'})
            
            if summary_data:
                summary_sheet = workbook.add_worksheet('Summary')
                row = 0
                summary_sheet.merge_range(row, 0, row, 3, '수입신고 분석 보고서', workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center'}))
                row += 2
                summary_sheet.write(row, 0, '전체 신고 건수', header_format)
                summary_sheet.write(row, 1, summary_data.get('전체 신고 건수', 0))
                row += 2
                
                for key in ['거래구분별', '세율구분별', 'Risk분석']:
                    if key in summary_data:
                        summary_sheet.write(row, 0, f'{key} 분석', header_format)
                        row += 1
                        summary_data[key].to_excel(writer, sheet_name='Summary', startrow=row, startcol=0, index=False)
                        row += len(summary_data[key]) + 2

            # 모든 결과 시트 저장
            sheet_map = {
                'eight_percent': '8% 환급 검토',
                'zero_risk': '0% 세율 위험',
                'tariff_risk': '세율 위험',
                'price_risk': '단가 위험',
                'domestic_tax': '내국세구분',
                'import_req_risk': '수입요건 Risk',
                'f_rate': 'F세율 적용',
                'fta_opp': 'FTA 기회 발굴',
                'low_price': '저가신고 의심',
                'currency_inc': '통화단위 불일치',
                'trade_type': '특수거래 구분',
                'free_freight': '무상운임 누락',
                'usage_rate': '용도세율 적용'
            }

            # 색상 포맷 정의
            header_blue_format = workbook.add_format({'bold': True, 'bg_color': '#4472C4', 'font_color': 'white', 'border': 1, 'align': 'center'})
            header_gray_format = workbook.add_format({'bold': True, 'bg_color': '#D9E1F2', 'border': 1, 'align': 'center'})
            red_format = workbook.add_format({'bg_color': '#FF0000', 'font_color': 'white', 'border': 1})
            yellow_format = workbook.add_format({'bg_color': '#FFFF00', 'border': 1})
            orange_format = workbook.add_format({'bg_color': '#FFA500', 'border': 1})
            text_format = workbook.add_format({'num_format': '@'})  # 텍스트 포맷

            for key, sheet_name in sheet_map.items():
                data = results.get(key)
                if data is not None and not data.empty:
                    # 데이터 먼저 쓰기
                    data.to_excel(writer, sheet_name=sheet_name, index=False)
                    worksheet = writer.sheets[sheet_name]
                    
                    # 세번부호와 수입신고번호 컬럼을 텍스트로 포맷 (소수점 방지)
                    text_columns = ['세번부호', '수입신고번호']
                    for col_name in text_columns:
                        if col_name in data.columns:
                            col_idx = data.columns.get_loc(col_name)
                            worksheet.set_column(col_idx, col_idx, None, text_format)
                    
                    # 중요 컬럼 정의 (파란색 헤더)
                    key_columns_map = {
                        'eight_percent': ['수입신고번호', '관세실행세율'],
                        'zero_risk': ['수입신고번호', '관세실행세율'],
                        'tariff_risk': ['규격1', '세번부호', '관세실행세율'],
                        'price_risk': ['수입신고번호', 'Z-Score'],
                        'domestic_tax': ['수입신고번호', '세번부호'],
                        'import_req_risk': ['규격1', '법령코드'],
                        'f_rate': ['수입신고번호', '세율구분'],
                        'fta_opp': ['수입신고번호', '관세실행세율'],
                        'low_price': ['수입신고번호', '단가'],
                        'currency_inc': ['무역거래처상호', '결제통화단위', '이상치점수'],
                        'free_freight': ['수입신고번호', '운임'],
                        'usage_rate': ['수입신고번호', '세율구분']
                    }
                    
                    key_cols = key_columns_map.get(key, [])
                    
                    # 헤더 색상 적용
                    for col_idx, col_name in enumerate(data.columns):
                        if col_name in key_cols:
                            worksheet.write(0, col_idx, col_name, header_blue_format)
                        else:
                            worksheet.write(0, col_idx, col_name, header_gray_format)
                    
                    # 데이터 행에 조건부 서식 적용
                    num_rows = len(data)
                    
                    # 각 시트별 특이값 색상 규칙 (빠른 조건부 서식만 사용)
                    if key == 'eight_percent' and '관세실행세율' in data.columns:
                        col_idx = data.columns.get_loc('관세실행세율')
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx, 
                            {'type': 'cell', 'criteria': '>=', 'value': 8, 'format': orange_format})
                    
                    elif key == 'zero_risk' and '관세실행세율' in data.columns:
                        col_idx = data.columns.get_loc('관세실행세율')
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx,
                            {'type': 'cell', 'criteria': '<', 'value': 8, 'format': red_format})
                    
                    elif key == 'fta_opp' and '관세실행세율' in data.columns:
                        col_idx = data.columns.get_loc('관세실행세율')
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx,
                            {'type': 'cell', 'criteria': '>', 'value': 0, 'format': orange_format})
                    
                    elif key == 'low_price' and '단가' in data.columns:
                        col_idx = data.columns.get_loc('단가')
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx,
                            {'type': 'cell', 'criteria': '<=', 'value': 10, 'format': yellow_format})
                    
                    elif key == 'currency_inc' and '이상치점수' in data.columns:
                        col_idx = data.columns.get_loc('이상치점수')
                        # 이상치점수 컬럼에 색상 적용
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx,
                            {'type': 'cell', 'criteria': '>', 'value': 80, 'format': yellow_format})
                        
                        # 이상치점수가 높은 행 전체를 강조 (행 단위 색상)
                        # 90 이상: 주황색 배경
                        if num_rows > 0:
                            for row_idx in range(1, min(num_rows + 1, 1001)):  # 최대 1000행까지만
                                try:
                                    score = data.iloc[row_idx - 1]['이상치점수']
                                    if pd.notna(score) and float(score) > 90:
                                        # 전체 행에 주황색 포맷 적용
                                        for col in range(len(data.columns)):
                                            cell_val = data.iloc[row_idx - 1, col]
                                            worksheet.write(row_idx, col, cell_val, orange_format)
                                except:
                                    pass
                    
                    elif key == 'free_freight' and '운임' in data.columns:
                        col_idx = data.columns.get_loc('운임')
                        worksheet.conditional_format(1, col_idx, num_rows, col_idx,
                            {'type': 'cell', 'criteria': '==', 'value': 0, 'format': yellow_format})
            
            

            # create_verification_methods_excel_sheet(writer)  # 속도 개선을 위해 제거
            
            
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"엑셀 생성 오류: {e}")
        return None

def create_word_document(results, summary_data):
    """워드 문서 생성 (특이건만 상세 포함)"""
    try:
        doc = Document()
        doc.add_heading('수입신고 RISK 분석 보고서', 0)
        doc.add_paragraph(datetime.datetime.now().strftime("%Y-%m-%d"))
        
        if summary_data:
            doc.add_heading('종합 요약', level=1)
            p = doc.add_paragraph()
            p.add_run(f"전체 신고 건수: {summary_data.get('전체 신고 건수', 0):,} 건").bold = True
            
            if 'Risk분석' in summary_data:
                risk_df = summary_data['Risk분석']
                risk_found = risk_df[risk_df['신고건수'] > 0]
                
                if len(risk_found) > 0:
                    p.add_run(f"\n\n⚠️ 발견된 Risk 유형: {len(risk_found)}건").bold = True
                    for _, row in risk_found.iterrows():
                        p.add_run(f"\n- {row['Risk 유형']}: {row['신고건수']:,} 건 ({row['비율(%)']:.1f}%)")
                else:
                    p.add_run("\n\n✅ 특이사항이 발견되지 않았습니다.").bold = True
        
        section_titles = {
            'eight_percent': ('8% 환급 검토', '8% 환급 검토  대상', ['수입신고번호', '세번부호', '관세실행세율', '금액', '거래품명']),
            'zero_risk': ('0% 세율 위험', '0% 세율 위험', ['수입신고번호', '세번부호', '세율구분', '관세실행세율', '거래품명']),
            'tariff_risk': ('세율 위험', '세율 위험(세번부호 불일치)', ['규격1', '세번부호', '세율구분', '거래품명']),
            'price_risk': ('단가 위험', '단가 이상치 (Z-Score)', ['수입신고번호', '규격1', '단가', 'Z-Score', '평균단가']),
            'domestic_tax': ('내국세구분 누락', '내국세구분 누락', ['수입신고번호', '세번부호', '거래품명', '금액']),
            'import_req_risk': ('수입요건 Risk', '수입요건 불일치', ['규격1', '수입신고번호', '법령코드', '발급서류명']),
            'f_rate': ('F세율 적용', 'F세율 적용 건', ['수입신고번호', '세번부호', '세율구분', '세율설명', '거래품명']),
            'fta_opp': ('FTA 기회 발굴', 'FTA 적용 기회', ['수입신고번호', '세번부호', '관세실행세율', '적출국코드', '원산지코드']),
            'low_price': ('저가신고 의심', '저가신고 의심 건', ['수입신고번호', '거래품명', '단가', '금액', '결제통화단위']),
            'currency_inc': ('통화단위 불일치', '통화단위 불일치 건', ['무역거래처상호', '결제통화단위', '수입신고번호', '금액']),
            'country_curr_inc': ('국가별 통화단위 불일치', '국가별 희귀 통화단위 사용', ['무역거래처국가코드', '결제통화단위', '사용비율', '이상치점수']),
            'trade_type': ('특수거래 구분', '특수거래 구분 건', ['수입신고번호', '거래구분', '세번부호', '거래품명', '금액']),
            'free_freight': ('무상운임 누락', '무상운임 누락 의심', ['수입신고번호', '결제방법', '운임', '금액', '거래품명']),
            'usage_rate': ('용도세율 적용', '용도세율 적용 건', ['수입신고번호', '세번부호', '세율구분', '세율설명', '거래품명'])
        }

        has_findings = False
        for key, (title, desc, display_cols) in section_titles.items():
            data =  results.get(key)
            if data is not None and not data.empty:
                has_findings = True
                doc.add_heading(title, level=1)
                doc.add_paragraph(f'총 {len(data):,} 건의 {desc}이(가) 식별되었습니다.')
                
                doc.add_paragraph('📋 상위 5건 샘플:', style='Heading 2')
                sample_data = data.head(5)
                sample_data = format_date_columns(sample_data)
                
                available_cols = [col for col in display_cols if col in sample_data.columns]
                if len(available_cols) == 0:
                    available_cols = sample_data.columns[:5].tolist()
                
                table = doc.add_table(rows=1, cols=len(available_cols))
                table.style = 'Light Grid Accent 1'
                
                header_cells = table.rows[0].cells
                for i, col_name in enumerate(available_cols):
                    header_cells[i].text = col_name
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                for _, row in sample_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, col_name in enumerate(available_cols):
                        value = row.get(col_name, '')
                        if isinstance(value, (int, float)) and not pd.isna(value):
                            if col_name in ['Z-Score', '평균단가', '표준편차', '사용비율', '이상치점수']:
                                row_cells[i].text = f"{value:.2f}"
                            else:
                                row_cells[i].text = f"{value:,.0f}" if value != 0 else "0"
                        else:
                            row_cells[i].text = str(value) if pd.notna(value) else ''
                
                doc.add_paragraph()
        
        if not has_findings:
            doc.add_heading('분석 결과', level=1)
            doc.add_paragraph('✅ 검토가 필요한 특이사항이 발견되지 않았습니다.')
        
        doc.add_paragraph()
        footer = doc.add_paragraph('Generated by 관세법인 우신')
        footer.alignment = 1
        
        doc_output = io.BytesIO()
        doc.save(doc_output)
        doc_output.seek(0)
        return doc_output.getvalue()
    except Exception as e:
        st.error(f"워드 문서 생성 중 오류 발생: {str(e)}")
        return None

def create_html_report(results, summary_data):
    """HTML 보고서 생성 (특이건만 상세 포함)"""
    try:
        html_content = f"""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>수입신고 RISK 분석 보고서</title>
    <style>
        body {{ font-family: 'Malgun Gothic', 'Segoe UI', sans-serif; margin: 0; padding: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 40px; border-radius: 10px; box-shadow: 0 10px 40px rgba(0,0,0,0.1); }}
        h1 {{ color: #2c3e50; text-align: center; font-size: 2.5em; margin-bottom: 10px; border-bottom: 3px solid #667eea; padding-bottom: 15px; }}
        .date {{ text-align: center; color: #7f8c8d; font-size: 1.1em; margin-bottom: 30px; }}
        .summary-box {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; margin: 30px 0; box-shadow: 0 5px 20px rgba(102, 126, 234, 0.3); }}
        .metric {{ display: inline-block; background: rgba(255,255,255,0.2); padding: 15px 25px; border-radius: 8px; margin: 10px; backdrop-filter: blur(10px); }}
        .metric-value {{ font-size: 2em; font-weight: bold; display: block; }}
        .section {{ margin: 40px 0; padding: 25px; background: #f8f9fa; border-radius: 8px; border-left: 5px solid #dc3545; }}
        .section h2 {{ color: #dc3545; margin-top: 0; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th {{ background: #667eea; color: white; padding: 12px; text-align: left; font-weight: bold; }}
        td {{ padding: 10px; border-bottom: 1px solid #ddd; }}
        tr:hover {{ background-color: #f5f5f5; }}
        .no-findings {{ text-align: center; color: #28a745; font-size: 1.3em; padding: 40px; }}
        .footer {{ text-align: center; margin-top: 50px; padding-top: 20px; border-top: 2px solid #ecf0f1; color: #7f8c8d; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📊 수입신고 RISK 분석 보고서</h1>
        <div class="date">{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}</div>
        
        <div class="summary-box">
            <h2>종합 요약</h2>
"""
        if summary_data:
            total_count = summary_data.get('전체 신고 건수', 0)
            html_content += f"""
            <div class="metric">
                <span class="metric-value">{total_count:,}</span>
                <span class="metric-label">전체 신고 건수</span>
            </div>
"""
            if 'Risk분석' in summary_data:
                risk_df = summary_data['Risk분석']
                risk_found = risk_df[risk_df['신고건수'] > 0]
                for _, row in risk_found.iterrows():
                    html_content += f"""
            <div class="metric">
                <span class="metric-value">{row['신고건수']:,}</span>
                <span class="metric-label">{row['Risk 유형']} ({row['비율(%)']:.1f}%)</span>
            </div>
"""
        html_content += "</div>"
        
        section_titles = {
            'eight_percent': '8% 환급 검토 대상',
            'zero_risk': '0% 세율 위험',
            'tariff_risk': '세율 위험(세번부호 불일치)',
            'price_risk': '단가 변동성 위험',
            'domestic_tax': '내국세구분 누락',
            'import_req_risk': '수입요건 불일치',
            'f_rate': 'F세율 적용',
            'fta_opp': 'FTA 적용 기회',
            'low_price': '저가신고 의심',
            'currency_inc': '통화단위 불일치',
            'country_curr_inc': '국가별 통화단위 불일치',
            'trade_type': '특수거래 구분',
            'free_freight': '무상운임 누락',
            'usage_rate': '용도세율 적용'
        }
        
        has_findings = False
        for key, desc in section_titles.items():
            data = results.get(key)
            if data is not None and not data.empty:
                has_findings = True
                html_content += f"""
        <div class="section">
            <h2>⚠️ {desc}</h2>
            <p>총 <strong>{len(data):,}</strong> 건의 {desc}이(가) 식별되었습니다.</p>
            <table>
                <thead>
                    <tr>
"""
                sample_data = format_date_columns(data.head(5))
                if len(sample_data) > 0:
                    cols_to_show = list(sample_data.columns[:6])
                    for col in cols_to_show:
                        html_content += f"<th>{col}</th>"
                    html_content += "</tr></thead><tbody>"
                    
                    for _, row in sample_data.iterrows():
                        html_content += "<tr>"
                        for col in cols_to_show:
                            value = row[col]
                            if isinstance(value, (int, float)) and not pd.isna(value):
                                html_content += f"<td>{value:,.2f}</td>"
                            else:
                                html_content += f"<td>{value}</td>"
                        html_content += "</tr>"
                    
                html_content += "</tbody></table></div>"
        
        if not has_findings:
            html_content += '<div class="no-findings">✅ 검토가 필요한 특이사항이 발견되지 않았습니다.</div>'
        
        html_content += """
        <div class="footer">
            <p><strong>Generated by 관세법인 우신</strong></p>
        </div>
    </div>
</body>
</html>
"""
        return html_content
    except Exception as e:
        st.error(f"HTML 보고서 생성 중 오류 발생: {str(e)}")
        return None

def main():
    col1, col2 = st.columns([1, 5])
    with col1:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(current_dir, 'logo.png')
        
        if os.path.exists(logo_path):
            st.image(logo_path, width=150)
        elif os.path.exists("logo.png"):
            st.image("logo.png", width=150)
    
    with col2:
        st.title("🛡️ TradeGuard (트레이드가드)")
        st.markdown("### 지능형 수입신고 리스크 분석 솔루션")
    
    st.markdown("---")
    
    st.sidebar.markdown("---")
    st.sidebar.caption("made by 전자동")

    uploaded_file = st.file_uploader("📁 엑셀 파일 업로드", type=['xlsx', 'xls', 'csv'])
    
    if uploaded_file is not None:
        progress_container = st.container()
        with progress_container:
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            df_original = read_excel_file(uploaded_file, progress_bar, status_text)
            
            if df_original is not None:
                time.sleep(0.5)
                progress_bar.empty()
                status_text.empty()
                
                st.success(f"📈 데이터 로드 완료: {len(df_original):,}건")
                
                with st.expander("📋 데이터 미리보기"):
                    st.dataframe(df_original.head(10).astype(str), use_container_width=True)
                
                st.sidebar.markdown("### 분석 옵션")
                
                all_options = [
                    "종합 분석", 
                    "8% 환급 검토", "0% 세율 위험", "세율 위험", "단가 위험", "내국세구분", "수입요건 Risk",
                    "F세율 적용", "FTA 기회 발굴", "저가신고 의심", "통화단위 불일치", "무상운임 누락", "용도세율 적용"
                ]
                
                analysis_options = st.sidebar.multiselect(
                    "수행할 분석을 선택하세요:",
                    all_options,
                    default=all_options
                )
                
                if st.sidebar.button("🔍 분석 시작", type="primary"):
                    results = {}
                    with st.spinner('분석 중...'):
                        if "종합 분석" in analysis_options: results['summary'] = create_summary_analysis(df_original)
                        if "8% 환급 검토" in analysis_options: results['eight_percent'] = create_eight_percent_refund_analysis(df_original)
                        if "0% 세율 위험" in analysis_options: results['zero_risk'] = create_zero_percent_risk_analysis(df_original)
                        if "세율 위험" in analysis_options: results['tariff_risk'] = create_tariff_risk_analysis(df_original)
                        if "단가 위험" in analysis_options: results['price_risk'] = create_price_risk_analysis(df_original)
                        if "내국세구분" in analysis_options: results['domestic_tax'] = create_domestic_tax_code_analysis(df_original)
                        if "수입요건 Risk" in analysis_options: results['import_req_risk'] = create_import_requirement_risk_analysis(df_original)
                        if "F세율 적용" in analysis_options: results['f_rate'] = create_f_rate_analysis(df_original)
                        if "FTA 기회 발굴" in analysis_options: results['fta_opp'] = create_fta_opportunity_analysis(df_original)
                        if "저가신고 의심" in analysis_options: results['low_price'] = create_low_price_analysis(df_original)
                        if "통화단위 불일치" in analysis_options: results['currency_inc'] = create_currency_consistency_analysis(df_original)
                        # "국가별 통화단위 불일치" 제거됨 (통화단위 불일치에 통합)
                        # "특수거래 구분" 제거됨 (사용자 요청)
                        if "무상운임 누락" in analysis_options: results['free_freight'] = create_free_charge_freight_analysis(df_original)
                        if "용도세율 적용" in analysis_options: results['usage_rate'] = create_usage_rate_analysis(df_original)
                    
                    st.success("분석 완료!")
                    
                    tabs = st.tabs([opt for opt in analysis_options if opt in all_options])
                    
                    key_map = {
                        "종합 분석": 'summary', 
                        "8% 환급 검토": 'eight_percent', 
                        "0% 세율 위험": 'zero_risk',
                        "세율 위험": 'tariff_risk', 
                        "단가 위험": 'price_risk', 
                        "내국세구분": 'domestic_tax',
                        "수입요건 Risk": 'import_req_risk',
                        "F세율 적용": 'f_rate',
                        "FTA 기회 발굴": 'fta_opp',
                        "저가신고 의심": 'low_price',
                        "통화단위 불일치": 'currency_inc',

                        "무상운임 누락": 'free_freight',
                        "용도세율 적용": 'usage_rate'
                    }

                    for i, tab_name in enumerate(tabs):
                       with tab_name:
                            key = key_map.get(analysis_options[i])
                            data = results.get(key)
                            
                            if key == 'summary' and data:
                                st.markdown("### 📈 종합 분석 대시보드")
                                m1, m2, m3, m4 = st.columns(4)
                                m1.metric("전체 신고 건수", f"{data.get('전체 신고 건수', 0):,}")
                                if 'Risk분석' in data:
                                    risk_df = data['Risk분석']
                                    for idx, row in risk_df.iterrows():
                                        if idx < 3:
                                            (m2 if idx==0 else m3 if idx==1 else m4).metric(
                                                row['Risk 유형'], 
                                                f"{row['신고건수']:,}", 
                                                f"{row['비율(%)']:.1f}%"
                                            )
                                st.markdown("---")
                                
                                c1, c2 = st.columns(2)
                                with c1:
                                    if 'Risk분석' in data:
                                        fig = px.pie(
                                            data['Risk분석'], 
                                            values='신고건수', 
                                            names='Risk 유형', 
                                            title='Risk 유형별 분포', 
                                            hole=0.4,
                                            color_discrete_sequence=px.colors.qualitative.Set3
                                        )
                                        st.plotly_chart(fig, use_container_width=True)
                                
                                with c2:
                                    if '월별추이' in data:
                                        monthly_df_display = data['월별추이'].copy()
                                        fig = px.line(
                                            monthly_df_display, 
                                            x='수리월', 
                                            y='신고건수', 
                                            title='월별 수입신고 추이', 
                                            markers=True
                                        )
                                        fig.update_xaxes(title_text='수리월 (년-월)', type='category')
                                        fig.update_layout(xaxis_tickangle=-45)
                                        st.plotly_chart(fig, use_container_width=True)
                                    else:
                                        st.info("월별 추이 데이터를 생성할 수 없습니다")

                            elif key == 'price_risk' and isinstance(data, pd.DataFrame) and not data.empty:
                                st.markdown("### 📊 단가 이상치 분포 (Z-Score 기준)")
                                
                                chart_data = data.copy()
                                chart_data[COL_ACCEPTANCE_DATE] = pd.to_numeric(chart_data[COL_ACCEPTANCE_DATE], errors='coerce').fillna(0).astype(int).astype(str)
                                chart_data[COL_ACCEPTANCE_DATE] = pd.to_datetime(chart_data[COL_ACCEPTANCE_DATE], format='%Y%m%d', errors='coerce')
                                
                                fig = px.scatter(
                                    chart_data, 
                                    x=COL_ACCEPTANCE_DATE, 
                                    y=COL_UNIT_PRICE,
                                    color=COL_SPEC_1,
                                    size=chart_data['Z-Score'].abs(),
                                    hover_data=[COL_TRADE_NAME, '평균단가', 'Z-Score'],
                                    title="이상치 산점도 (점 크기: Z-Score 절대값)"
                                )
                                st.plotly_chart(fig, use_container_width=True)
                                
                                display_df = format_date_columns(data)
                                st.dataframe(display_df.astype(str), use_container_width=True)

                            elif isinstance(data, pd.DataFrame) and not data.empty:
                                display_df = format_date_columns(data)
                                st.dataframe(display_df.astype(str), use_container_width=True)
                            else:
                                st.info("해당하는 데이터가 없습니다.")

                    st.markdown("---")
                    st.subheader("📥 결과 다운로드")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        excel_data = create_excel_file(df_original, results, results.get('summary', {}))
                        if excel_data:
                            st.download_button("📊 엑셀 보고서", excel_data, f"수입신고분석_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                            
                    with col2:
                        word_data = create_word_document(results, results.get('summary', {}))
                        if word_data:
                            st.download_button("📄 워드 보고서", word_data, f"수입신고분석_{datetime.datetime.now().strftime('%Y%m%d')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                            
                    with col3:
                        html_data = create_html_report(results, results.get('summary', {}))
                        if html_data:
                            st.download_button("🌐 HTML 보고서", html_data, f"수입신고분석_{datetime.datetime.now().strftime('%Y%m%d')}.html", "text/html", use_container_width=True)

if __name__ == "__main__":
    main()
